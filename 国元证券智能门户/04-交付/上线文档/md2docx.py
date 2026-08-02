#!/usr/bin/env python3
"""Markdown → Word 转换器，中文排版，10份文档批处理"""

import re
import os
import glob
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ── 样式配置 ──────────────────────────────
FONT_BODY = '宋体'
FONT_BODY_WEST = 'Times New Roman'
FONT_HEADING = '黑体'
FONT_HEADING_WEST = 'Arial'
FONT_CODE = 'Consolas'
SIZE_BODY = Pt(10.5)
SIZE_H1 = Pt(18)
SIZE_H2 = Pt(15)
SIZE_H3 = Pt(13)
SIZE_H4 = Pt(12)
SIZE_CODE = Pt(8.5)
SIZE_FOOTER = Pt(8)
PAGE_A4 = (Cm(21), Cm(29.7))
MARGIN = Cm(2.54)

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))
WORD_DIR = os.path.join(OUTPUT_DIR, 'Word版')
os.makedirs(WORD_DIR, exist_ok=True)

def set_font(run, font_name=FONT_BODY, font_west=FONT_BODY_WEST, size=SIZE_BODY, bold=False, color=None):
    """设置run的字体"""
    run.font.size = size
    run.bold = bold
    run.font.name = font_name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), font_name)
    rFonts.set(qn('w:ascii'), font_west)
    rFonts.set(qn('w:hAnsi'), font_west)
    if color:
        run.font.color.rgb = color

def add_paragraph(doc, text='', style=None, font_name=FONT_BODY, font_west=FONT_BODY_WEST, font_size=SIZE_BODY, bold=False, alignment=None, space_after=Pt(4)):
    """添加段落并设置格式"""
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_font(run, font_name, font_west, font_size, bold)
    if alignment is not None:
        p.alignment = alignment
    pf = p.paragraph_format
    pf.space_after = space_after
    pf.line_spacing = 1.5
    return p

def add_heading_with_style(doc, text, level):
    """根据级别添加标题"""
    sizes = {1: SIZE_H1, 2: SIZE_H2, 3: SIZE_H3, 4: SIZE_H4}
    size = sizes.get(level, SIZE_H4)
    p = add_paragraph(doc, text, font_name=FONT_HEADING, font_west=FONT_HEADING_WEST, font_size=size, bold=True, space_after=Pt(8))
    pf = p.paragraph_format
    pf.space_before = Pt(12) if level <= 2 else Pt(8)
    return p

def add_code_block(doc, lines):
    """添加代码块"""
    for line in lines:
        p = add_paragraph(doc, line, font_name=FONT_CODE, font_west=FONT_CODE, font_size=SIZE_CODE, space_after=Pt(0))
        pf = p.paragraph_format
        pf.space_before = Pt(0)
        pf.line_spacing = 1.2
        # 添加灰色背景效果（通过段落底纹）
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
        p._element.get_or_add_pPr().append(shading)

def parse_inline(doc, text, base_run=None):
    """解析行内格式：粗体、斜体、行内代码"""
    # 这个方法返回一个 run 列表
    runs = []
    i = 0
    current = ''
    while i < len(text):
        # 行内代码 `...`
        if text[i] == '`' and i+1 < len(text):
            j = text.index('`', i+1) if '`' in text[i+1:] else -1
            if j > i:
                if current:
                    runs.append(('normal', current))
                    current = ''
                runs.append(('code', text[i+1:j]))
                i = j+1
                continue
        # 粗体 **...**
        if text[i:i+2] == '**' and i+2 < len(text):
            j = text.find('**', i+2)
            if j > i:
                if current:
                    runs.append(('normal', current))
                    current = ''
                inner = text[i+2:j]
                runs.append(('bold', inner))
                i = j+2
                continue
        # 斜体 *...* (单星号，但不要匹配 **)
        if text[i] == '*' and i+1 < len(text) and text[i+1] != '*':
            j = text.find('*', i+1)
            if j > i:
                if current:
                    runs.append(('normal', current))
                    current = ''
                runs.append(('italic', text[i+1:j]))
                i = j+1
                continue
        current += text[i]
        i += 1
    if current:
        runs.append(('normal', current))
    return runs

def process_markdown_to_docx(md_path, docx_path):
    """将Markdown文件转换为Word文档"""
    doc = Document()
    
    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = PAGE_A4[0]
    section.page_height = PAGE_A4[1]
    section.top_margin = MARGIN
    section.bottom_margin = MARGIN
    section.left_margin = MARGIN
    section.right_margin = MARGIN
    
    # ── 页脚（页码） ──
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 使用 Word 页码域
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run = fp.add_run()
    set_font(run, size=SIZE_FOOTER, color=RGBColor(0x99, 0x99, 0x99))
    run._element.append(fldChar1)
    run2 = fp.add_run()
    set_font(run2, size=SIZE_FOOTER, color=RGBColor(0x99, 0x99, 0x99))
    run2._element.append(instrText)
    run3 = fp.add_run()
    set_font(run3, size=SIZE_FOOTER, color=RGBColor(0x99, 0x99, 0x99))
    run3._element.append(fldChar2)
    
    # ── 读取并解析 ──
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []
    
    while i < len(lines):
        line = lines[i].rstrip('\n')
        
        # 代码块
        if line.strip().startswith('```'):
            if in_code_block:
                if code_lines:
                    add_code_block(doc, code_lines)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # 空行
        if not line.strip():
            if in_table and table_rows:
                # 渲染表格
                _render_table(doc, table_rows)
                table_rows = []
                in_table = False
            i += 1
            continue
        
        # 标题
        heading_match = re.match(r'^(#{1,4})\s+(.+)', line)
        if heading_match:
            if in_table and table_rows:
                _render_table(doc, table_rows)
                table_rows = []
                in_table = False
            level = len(heading_match.group(1))
            title = heading_match.group(2).strip()
            # 去除行内粗体标记
            title = re.sub(r'\*\*(.+?)\*\*', r'\1', title)
            add_heading_with_style(doc, title, level)
            i += 1
            continue
        
        # 表格
        if line.strip().startswith('|') and line.strip().endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            # 跳过分隔行
            if re.match(r'^\|[\s\-:|]+\|$', line.strip()):
                i += 1
                continue
            cells = [c.strip() for c in line.strip().split('|')[1:-1]]
            # 去除行内格式标记
            clean_cells = []
            for c in cells:
                c = re.sub(r'\*\*(.+?)\*\*', r'\1', c)
                c = re.sub(r'\*(.+?)\*', r'\1', c)
                c = re.sub(r'`(.+?)`', r'\1', c)
                clean_cells.append(c)
            table_rows.append(clean_cells)
            i += 1
            continue
        
        # 普通行（表格外的）
        if in_table and table_rows:
            _render_table(doc, table_rows)
            table_rows = []
            in_table = False
        
        # 无序列表
        ul_match = re.match(r'^(\s*)[-*+]\s+(.+)', line)
        if ul_match:
            content = ul_match.group(2)
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
            content = re.sub(r'`(.+?)`', r'\1', content)
            p = add_paragraph(doc, '• ' + content, font_size=SIZE_BODY, space_after=Pt(2))
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            i += 1
            continue
        
        # 有序列表
        ol_match = re.match(r'^(\s*)\d+[\.\)]\s+(.+)', line)
        if ol_match:
            content = ol_match.group(2)
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
            content = re.sub(r'`(.+?)`', r'\1', content)
            p = add_paragraph(doc, content, font_size=SIZE_BODY, space_after=Pt(2))
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            i += 1
            continue
        
        # 引用块
        quote_match = re.match(r'^>\s*(.+)', line)
        if quote_match:
            content = quote_match.group(1)
            content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
            p = add_paragraph(doc, content, font_size=Pt(9.5), space_after=Pt(4))
            pf = p.paragraph_format
            pf.left_indent = Cm(1)
            # 添加左边框效果（灰色底纹）
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8F8F8" w:val="clear"/>')
            p._element.get_or_add_pPr().append(shading)
            i += 1
            continue
        
        # 分隔线
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            p = add_paragraph(doc, '─' * 40, font_size=Pt(6), space_after=Pt(4), alignment=WD_ALIGN_PARAGRAPH.CENTER)
            i += 1
            continue
        
        # 普通段落
        content = line.strip()
        if content:
            # 检查是否有行内格式
            if '**' in content or '`' in content:
                # 去掉行内标记，保留纯文本
                content = re.sub(r'\*\*(.+?)\*\*', r'\1', content)
                content = re.sub(r'\*(.+?)\*', r'\1', content)
                content = re.sub(r'`(.+?)`', r'\1', content)
            add_paragraph(doc, content, font_size=SIZE_BODY, space_after=Pt(4))
        
        i += 1
    
    # 处理末尾表格
    if in_table and table_rows:
        _render_table(doc, table_rows)
    
    # ── 保存 ──
    doc.save(docx_path)
    return True

def _render_table(doc, rows):
    """渲染表格"""
    if not rows:
        return
    
    max_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci < max_cols:
                cell = table.cell(ri, ci)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                # 表头加粗
                if ri == 0:
                    set_font(run, font_name=FONT_HEADING, font_west=FONT_HEADING_WEST, size=Pt(9), bold=True)
                    # 表头背景色
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D6E4F0" w:val="clear"/>')
                    cell._element.get_or_add_tcPr().append(shading)
                else:
                    set_font(run, size=Pt(9))
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 表格后空行
    add_paragraph(doc, '', space_after=Pt(4))

# ── 主程序 ──
if __name__ == '__main__':
    md_files = sorted(glob.glob(os.path.join(OUTPUT_DIR, 'D*.md')))
    total = len(md_files)
    
    print(f'找到 {total} 份文档，开始转换...\n')
    
    for md_file in md_files:
        basename = os.path.splitext(os.path.basename(md_file))[0]
        docx_file = os.path.join(WORD_DIR, f'{basename}.docx')
        
        try:
            process_markdown_to_docx(md_file, docx_file)
            size_kb = os.path.getsize(docx_file) / 1024
            print(f'  ✅ {basename:40s} → {size_kb:.0f} KB')
        except Exception as e:
            print(f'  ❌ {basename:40s} → {e}')
    
    print(f'\n全部完成！输出目录: {WORD_DIR}')
